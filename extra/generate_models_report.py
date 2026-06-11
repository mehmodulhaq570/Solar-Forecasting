"""Generate Word report: All 6 Model Explanations for Solar FYP"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
section = doc.sections[0]
section.top_margin    = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin   = Inches(1.25)
section.right_margin  = Inches(1.25)

# ── Helpers ───────────────────────────────────────────────────────────────────
def heading1(text):
    p = doc.add_heading(text, level=1)
    p.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    return p

def heading2(text):
    p = doc.add_heading(text, level=2)
    p.runs[0].font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
    return p

def heading3(text):
    p = doc.add_heading(text, level=3)
    p.runs[0].font.color.rgb = RGBColor(0x40, 0x40, 0x40)
    return p

def body(text):
    p = doc.add_paragraph(text)
    for run in p.runs:
        run.font.size = Pt(11)
    return p

def bullet(text):
    p = doc.add_paragraph(text, style="List Bullet")
    for run in p.runs:
        run.font.size = Pt(11)
    return p

def bullet2(text):
    p = doc.add_paragraph(text, style="List Bullet 2")
    for run in p.runs:
        run.font.size = Pt(10)
    return p

def code_block(text):
    p = doc.add_paragraph(text)
    p.style = doc.styles["Normal"]
    for run in p.runs:
        run.font.name = "Courier New"
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x00, 0x00, 0x80)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F0F0F0')
    pPr.append(shd)
    return p

def add_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        run = hdr[i].paragraphs[0].runs[0]
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        tc = hdr[i]._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), '1F497D')
        tcPr.append(shd)
    for ri, row_data in enumerate(rows):
        cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row_data):
            cells[ci].text = str(val)
            cells[ci].paragraphs[0].runs[0].font.size = Pt(10)
            if ri % 2 == 1:
                tc = cells[ci]._tc
                tcPr = tc.get_or_add_tcPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:color'), 'auto')
                shd.set(qn('w:fill'), 'DCE6F1')
                tcPr.append(shd)
    if col_widths:
        for row in table.rows:
            for ci, cell in enumerate(row.cells):
                cell.width = Inches(col_widths[ci])
    return table

def centered_formula(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    return p

def performance_box(train_r2, test_r2, rmse, mae):
    add_table(
        ["Metric", "Train", "Test"],
        [
            ["R²",           train_r2, test_r2],
            ["RMSE (W/m²)",  "—",      rmse],
            ["MAE (W/m²)",   "—",      mae],
        ],
        col_widths=[1.6, 1.4, 1.4],
    )

# ═══════════════════════════════════════════════════════════════════════════════
# TITLE
# ═══════════════════════════════════════════════════════════════════════════════
title = doc.add_heading("Solar Radiation Forecasting — Model Explanations", 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

sub = doc.add_paragraph("Detailed Architecture, Training Strategy & Performance of All 6 Models")
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.runs[0].bold = True
sub.runs[0].font.size = Pt(12)
sub.runs[0].font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
doc.add_paragraph("")

# ── Overview table ─────────────────────────────────────────────────────────────
heading1("Overview: All 6 Models at a Glance")
add_table(
    ["#", "Model", "Framework", "Type", "Input Format", "Test R²", "Test RMSE"],
    [
        ["1", "XGBoost",      "xgboost",          "Gradient Boosted Trees",        "41 flat features per hour",     "0.9933", "22.62"],
        ["2", "Random Forest","scikit-learn",      "Bagged Decision Trees",         "41 flat features per hour",     "0.9933", "22.77"],
        ["3", "LSTM",         "TensorFlow/Keras",  "Bidirectional Recurrent NN",    "(N, 24, 15) sequences",         "0.9873", "31.30"],
        ["4", "CNN-LSTM",     "TensorFlow/Keras",  "Conv + Recurrent Hybrid",       "(N, 24, 15) sequences",         "0.9896", "28.22"],
        ["5", "TFT",          "PyTorch Lightning", "Transformer-based",             "(N, 24, 14) sequences",         "0.9472", "63.69"],
        ["6", "TCN",          "PyTorch Lightning", "Dilated Causal Convolutions",   "(N, 24, 14) sequences",         "0.9461", "64.38"],
    ],
    col_widths=[0.3, 1.3, 1.5, 1.8, 1.8, 0.8, 0.8],
)
doc.add_paragraph("")
body("All models predict SolarRadiation (W/m²) for each hour of a target day using features derived "
     "from NASA POWER historical meteorological data for Lahore, Pakistan.")

# ═══════════════════════════════════════════════════════════════════════════════
# MODEL 1 — XGBOOST
# ═══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
heading1("Model 1: XGBoost (Extreme Gradient Boosting)")

heading2("1.1 What Is XGBoost?")
body("XGBoost is an ensemble learning algorithm based on gradient boosting. It builds a sequence "
     "of decision trees where each new tree corrects the residual errors of the previous trees. "
     "'Extreme' refers to the system-level and algorithmic optimizations (parallelism, cache-aware "
     "computation, sparsity handling) that make it much faster and more regularized than standard GBM.")

heading2("1.2 How It Works (Conceptually)")
body("At each boosting round t, a new tree f_t is added to minimize the residual error:")
centered_formula("F_t(x)  =  F_{t-1}(x)  +  η · f_t(x)")
body("where η is the learning rate (shrinkage). The tree f_t is fit to the negative gradient of the "
     "loss function (MSE in this case), effectively doing gradient descent in function space.")
bullet("Each tree is a shallow regression tree (max_depth=5).")
bullet("At each split, only a random subset of features (colsample_bytree=0.8) and rows (subsample=0.8) "
       "are considered, adding stochasticity to reduce overfitting.")
bullet("The final prediction is the sum of 800 trees, each contributing a small improvement.")

heading2("1.3 Architecture / Hyperparameters")
add_table(
    ["Hyperparameter", "Value", "Rationale"],
    [
        ["n_estimators",      "800",   "More trees = finer correction; EarlyStopping not used here so 800 is the total budget"],
        ["max_depth",         "5",     "Shallow trees reduce overfitting and keep individual trees interpretable"],
        ["learning_rate (η)", "0.05",  "Small step size requires more trees but leads to better generalization"],
        ["subsample",         "0.8",   "Row sampling — each tree sees 80% of training data (stochastic GBM)"],
        ["colsample_bytree",  "0.8",   "Feature sampling — each tree uses 80% of features"],
        ["random_state",      "42",    "Reproducibility"],
        ["n_jobs",            "-1",    "Use all CPU cores for parallel tree building"],
        ["Objective",         "reg:squarederror (MSE)", "Regression with squared error loss"],
    ],
    col_widths=[1.8, 1.2, 3.3],
)

heading2("1.4 Input Features")
body("XGBoost receives a flat feature vector of 41 values per prediction step (1 hour):")
bullet("17 base features: 3 raw time, 6 cyclical encodings, 2 solar geometry, 6 weather variables.")
bullet("24 lag features: SolarRadiation at t-1, t-2, …, t-24 (autoregressive memory).")
bullet("During inference, lag features are updated autoregressively — each predicted hour's output "
       "becomes the lag_1 for the next hour.")

heading2("1.5 Training Process")
bullet("Fit on training set (80% chronological) in one pass — no epochs, no batches.")
bullet("No validation split required; XGBoost trees are non-parametric and do not need gradient descent.")
bullet("Saved to: xgboost_model.pkl via joblib.")

heading2("1.6 Performance")
performance_box("0.9981", "0.9933", "22.62", "9.69")
doc.add_paragraph("")
body("XGBoost is the top-performing individual model. The tiny train/test R² gap (0.9981 → 0.9933) "
     "shows minimal overfitting, which is impressive given 41 features. This is attributed to "
     "feature subsampling (colsample_bytree=0.8) and row subsampling (subsample=0.8).")

# ═══════════════════════════════════════════════════════════════════════════════
# MODEL 2 — RANDOM FOREST
# ═══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
heading1("Model 2: Random Forest")

heading2("2.1 What Is Random Forest?")
body("Random Forest is a bagging (Bootstrap Aggregating) ensemble of decision trees. Unlike boosting "
     "(which is sequential), bagging trains trees independently in parallel on different bootstrap "
     "samples of the data. Final prediction is the average (mean) of all tree predictions.")
body("The key randomness comes from two sources:")
bullet("Bootstrap sampling: Each tree is trained on a random sample (with replacement) of the training data.")
bullet("Feature randomness: At each node split, only a random subset of √(n_features) features is considered.")

heading2("2.2 How It Works (Conceptually)")
centered_formula("F(x)  =  (1/T) · Σ_{t=1}^{T}  f_t(x)")
body("where T=300 independent trees each trained on a bootstrap sample. The variance reduction from "
     "averaging 300 trees makes Random Forest highly robust to overfitting compared to a single deep tree.")

heading2("2.3 Architecture / Hyperparameters")
add_table(
    ["Hyperparameter", "Value", "Rationale"],
    [
        ["n_estimators",  "300",    "300 trees gives stable averaging; diminishing returns beyond this"],
        ["max_depth",     "10",     "Limits each tree depth to prevent overfitting individual trees"],
        ["random_state",  "42",     "Reproducibility"],
        ["n_jobs",        "-1",     "Fully parallel tree training across all CPU cores"],
        ["Criterion",     "squared_error (MSE)", "Standard regression split criterion"],
        ["max_features",  "auto (√n)",  "Random feature subset at each split — key to diversity between trees"],
    ],
    col_widths=[1.8, 1.5, 3.0],
)

heading2("2.4 Input Features")
body("Identical to XGBoost: 17 base features + 24 lag features = 41 total.")
body("Random Forest and XGBoost share the exact same feature matrix (X_train_tree, X_test_tree), "
     "making their outputs directly comparable.")

heading2("2.5 Training Process")
bullet("One-pass training on all 300 trees in parallel — no epochs.")
bullet("No scaling required: decision trees are invariant to feature scale.")
bullet("Saved to: random_forest_model.pkl via joblib.")

heading2("2.6 XGBoost vs. Random Forest — Key Differences")
add_table(
    ["Aspect", "XGBoost", "Random Forest"],
    [
        ["Training strategy",  "Sequential (boosting)",             "Parallel (bagging)"],
        ["Tree count",         "800",                               "300"],
        ["Tree depth",         "5 (shallow)",                       "10 (deeper)"],
        ["Bias-Variance",      "Low bias via boosting corrections", "Low variance via averaging"],
        ["Speed",              "Slower (sequential)",              "Faster (parallel)"],
        ["Test R²",            "0.9933",                            "0.9933"],
        ["Test RMSE",          "22.62 W/m²",                        "22.77 W/m²"],
    ],
    col_widths=[1.8, 2.5, 2.5],
)
doc.add_paragraph("")
body("Both achieve identical Test R² = 0.9933. XGBoost has marginally better RMSE (22.62 vs 22.77) "
     "with more trees (800 vs 300) due to sequential refinement.")

heading2("2.7 Performance")
performance_box("0.9963", "0.9933", "22.77", "9.88")

# ═══════════════════════════════════════════════════════════════════════════════
# MODEL 3 — LSTM
# ═══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
heading1("Model 3: LSTM (Bidirectional Long Short-Term Memory)")

heading2("3.1 What Is LSTM?")
body("LSTM (Long Short-Term Memory) is a type of Recurrent Neural Network (RNN) designed to learn "
     "long-range temporal dependencies. Standard RNNs suffer from the vanishing gradient problem — "
     "gradients shrink exponentially as they propagate back through time, making it hard to learn "
     "patterns separated by many timesteps. LSTMs solve this with a gated memory cell.")

heading2("3.2 LSTM Cell Mechanics")
body("Each LSTM cell maintains two state vectors: hidden state h_t (short-term) and cell state c_t (long-term).")
body("Four gates control information flow:")
add_table(
    ["Gate", "Formula", "Purpose"],
    [
        ["Forget gate",  "f_t = σ(W_f · [h_{t-1}, x_t] + b_f)",  "Decides what fraction of old memory c_{t-1} to erase"],
        ["Input gate",   "i_t = σ(W_i · [h_{t-1}, x_t] + b_i)",  "Controls how much new information to write"],
        ["Cell update",  "c̃_t = tanh(W_c · [h_{t-1}, x_t] + b_c)", "Candidate new memory content"],
        ["Cell state",   "c_t = f_t ⊙ c_{t-1}  +  i_t ⊙ c̃_t",  "Updated long-term memory"],
        ["Output gate",  "o_t = σ(W_o · [h_{t-1}, x_t] + b_o)",  "What fraction of cell state to expose as output"],
        ["Hidden state", "h_t = o_t ⊙ tanh(c_t)",                 "Short-term output used as input to next step"],
    ],
    col_widths=[1.4, 2.8, 2.1],
)

heading2("3.3 Bidirectional Extension")
body("A standard LSTM only processes the sequence forward (past → future). A Bidirectional LSTM runs "
     "two LSTMs in parallel — one forward, one backward — and concatenates their hidden states:")
centered_formula("h_t  =  [h_t_forward  ||  h_t_backward]")
body("For solar radiation, this allows the model to use both past context (e.g., yesterday's radiation) "
     "AND future context within the 24-step window (e.g., if hour 20 shows high radiation, it informs "
     "hour 18 even though it comes later in the sequence).")

heading2("3.4 Architecture (Layer-by-Layer)")
add_table(
    ["Layer", "Type", "Parameters", "Output Shape"],
    [
        ["Input",   "Input",                      "shape=(24, 15)",          "(batch, 24, 15)"],
        ["LSTM-1",  "Bidirectional(LSTM(64))",     "units=64, return_seq=True",  "(batch, 24, 128)"],
        ["Drop-1",  "Dropout",                     "rate=0.2",               "(batch, 24, 128)"],
        ["LSTM-2",  "Bidirectional(LSTM(32))",     "units=32, return_seq=False", "(batch, 64)"],
        ["Drop-2",  "Dropout",                     "rate=0.2",               "(batch, 64)"],
        ["Dense-1", "Dense + ReLU",                "units=64",               "(batch, 64)"],
        ["Dense-2", "Dense + ReLU",                "units=32",               "(batch, 32)"],
        ["Output",  "Dense (linear)",              "units=1",                "(batch, 1)"],
    ],
    col_widths=[1.0, 2.0, 2.0, 1.8],
)
doc.add_paragraph("")
body("The first BiLSTM(64) returns the full sequence (return_sequences=True) so the second BiLSTM(32) "
     "can also process the temporal dimension. The second BiLSTM returns only the final timestep, "
     "collapsing the sequence into a single feature vector for the Dense head.")

heading2("3.5 Training Configuration")
add_table(
    ["Setting", "Value"],
    [
        ["Loss function",          "MSE (Mean Squared Error)"],
        ["Optimizer",              "Adam, learning_rate=0.001"],
        ["Batch size",             "32"],
        ["Max epochs",             "150"],
        ["Validation split",       "10% of training data (last 10%, chronological)"],
        ["Early stopping",         "patience=15, monitor=val_loss, restore_best_weights=True"],
        ["LR scheduler",           "ReduceLROnPlateau: factor=0.5, patience=5, min_lr=1e-6"],
        ["Input normalization",    "StandardScaler fit on training sequences (scaler_seq)"],
        ["Output scaling",         "Separate scaler_y for inverse_transform of predictions"],
        ["Saved as",               "lstm_model.keras (v3) / lstm_model.h5 (legacy)"],
    ],
    col_widths=[2.0, 4.3],
)

heading2("3.6 Performance")
performance_box("0.9876", "0.9873", "31.30", "20.19")
doc.add_paragraph("")
body("The LSTM's higher MAE (20.19 vs XGBoost's 9.69) reflects the inherent difficulty recurrent "
     "models face with sharp intra-day transitions (dawn/dusk ramps). Its advantage is in learning "
     "smooth multi-day seasonal trends that tree models may miss.")

# ═══════════════════════════════════════════════════════════════════════════════
# MODEL 4 — CNN-LSTM
# ═══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
heading1("Model 4: CNN-LSTM (Convolutional + Bidirectional LSTM)")

heading2("4.1 Motivation: Why Combine CNN and LSTM?")
body("Pure LSTM processes the sequence step-by-step, treating each timestep independently before "
     "passing it to the next. This can miss local short-range patterns (e.g., a 2–3 hour ramp in "
     "radiation at sunrise) that a convolutional filter can extract in one operation.")
body("CNN-LSTM uses Conv1D layers as a local feature extractor before the LSTM layers:")
bullet("Conv1D captures local temporal motifs (e.g., sharp dawn increase, cloud pass-through dips).")
bullet("MaxPool1D reduces the sequence length, forcing the LSTM to work on abstract higher-level features.")
bullet("BiLSTM then models the longer-range dependencies across the condensed sequence.")

heading2("4.2 Architecture (Layer-by-Layer)")
add_table(
    ["Layer", "Type", "Parameters", "Output Shape"],
    [
        ["Input",     "Input",             "shape=(24, 15)",                 "(batch, 24, 15)"],
        ["Conv-1",    "Conv1D + ReLU",     "filters=64, kernel=3, padding=same", "(batch, 24, 64)"],
        ["Conv-2",    "Conv1D + ReLU",     "filters=64, kernel=3, padding=same", "(batch, 24, 64)"],
        ["MaxPool",   "MaxPool1D",         "pool_size=2",                    "(batch, 12, 64)"],
        ["Drop-1",    "Dropout",           "rate=0.2",                       "(batch, 12, 64)"],
        ["BiLSTM-1",  "Bidirectional(LSTM(64))", "units=64, return_seq=True", "(batch, 12, 128)"],
        ["Drop-2",    "Dropout",           "rate=0.2",                       "(batch, 12, 128)"],
        ["BiLSTM-2",  "Bidirectional(LSTM(32))", "units=32, return_seq=False", "(batch, 64)"],
        ["Drop-3",    "Dropout",           "rate=0.2",                       "(batch, 64)"],
        ["Dense-1",   "Dense + ReLU",      "units=64",                       "(batch, 64)"],
        ["Dense-2",   "Dense + ReLU",      "units=32",                       "(batch, 32)"],
        ["Output",    "Dense (linear)",    "units=1",                        "(batch, 1)"],
    ],
    col_widths=[0.9, 2.0, 2.2, 1.7],
)

heading2("4.3 How Conv1D Filters Work on Time Series")
body("Each Conv1D filter slides a learnable kernel of size 3 across the 24-timestep sequence. "
     "With 64 filters and kernel_size=3, the layer learns 64 different 3-hour temporal patterns "
     "(e.g., a rising edge, a plateau, a falling edge). padding='same' preserves the sequence length.")
body("MaxPool1D(2) then takes the maximum value in non-overlapping windows of size 2, halving the "
     "sequence from 24 → 12 steps. This provides translational invariance (the pattern is detected "
     "regardless of exact position) and reduces the input length for the LSTM.")

heading2("4.4 CNN-LSTM vs. Plain LSTM — Key Differences")
add_table(
    ["Aspect", "LSTM", "CNN-LSTM"],
    [
        ["Local pattern detection",   "Weak (sequential step processing)",  "Strong (Conv1D kernels)"],
        ["Sequence length to LSTM",   "24 steps",                           "12 steps (after MaxPool)"],
        ["Overfitting risk",          "Moderate",                           "Lower (Conv regularizes)"],
        ["Test R²",                   "0.9873",                             "0.9896"],
        ["Test RMSE",                 "31.30 W/m²",                         "28.22 W/m²"],
        ["Test MAE",                  "20.19 W/m²",                         "17.01 W/m²"],
        ["Parameter count",           "~Fewer",                             "~More (Conv weights added)"],
    ],
    col_widths=[2.0, 2.0, 2.8],
)

heading2("4.5 Training Configuration")
body("Identical to LSTM: Adam(lr=0.001), MSE loss, batch=32, max_epochs=150, "
     "EarlyStopping(patience=15), ReduceLROnPlateau(patience=5). "
     "Saved as: cnn_lstm_model.keras / cnn_lstm_model.h5.")

heading2("4.6 Performance")
performance_box("0.9891", "0.9896", "28.22", "17.01")
doc.add_paragraph("")
body("CNN-LSTM outperforms plain LSTM on all metrics. The ~3 W/m² RMSE improvement and ~3 W/m² MAE "
     "improvement confirm that local convolutional feature extraction is beneficial for solar "
     "radiation patterns which have strong short-range structure (sunrise/sunset ramps).")

# ═══════════════════════════════════════════════════════════════════════════════
# MODEL 5 — TFT
# ═══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
heading1("Model 5: TFT (Temporal Fusion Transformer)")

heading2("5.1 What Is the Temporal Fusion Transformer?")
body("The Temporal Fusion Transformer (TFT) is a deep learning architecture based on the "
     "self-attention mechanism introduced in the original Transformer (Vaswani et al., 2017). "
     "Unlike RNNs that process sequences step-by-step, Transformer encoders process all timesteps "
     "simultaneously via self-attention, making them better at capturing global long-range dependencies.")
body("In this project, a simplified TFT is implemented with the core Transformer encoder block, "
     "learnable positional encoding, and a regression head.")

heading2("5.2 Self-Attention Mechanism")
body("The key operation in each Transformer layer is Multi-Head Self-Attention (MHSA):")
centered_formula("Attention(Q, K, V)  =  softmax( QK^T / √d_k ) · V")
body("Where Q (Queries), K (Keys), and V (Values) are all linear projections of the input sequence. "
     "Each timestep attends to all other timesteps, learning which past/future hours are most "
     "relevant for predicting radiation at any given step.")
body("Multi-head attention runs this operation in parallel with num_heads=4 different projection "
     "matrices, then concatenates and linearly projects the results. This allows the model to "
     "attend to different aspects of the sequence simultaneously.")

heading2("5.3 Architecture (Layer-by-Layer)")
add_table(
    ["Component", "Description", "Parameters"],
    [
        ["Input Projection",      "Linear(input_size → hidden_size=64)",                 "Projects 14 features to 64-dim space"],
        ["Positional Encoding",   "Learnable parameter: (1, 100, 64)",                   "Adds position information (unlike fixed sinusoidal PE)"],
        ["Transformer Layer 1",   "TransformerEncoderLayer(d_model=64, nhead=4, ffn=256, activation=GELU)", "MHSA + Feed-Forward sublayers, pre-norm"],
        ["Transformer Layer 2",   "Same as Layer 1",                                     "Second encoder layer for deeper representation"],
        ["Last Timestep Select",  "encoded_output[:, -1, :]",                            "Extracts final timestep representation for prediction"],
        ["LayerNorm",             "LayerNorm(64)",                                        "Stabilizes output before dense layers"],
        ["Dense-1",               "Linear(64 → 32) + GELU + Dropout(0.2)",              "First projection in regression head"],
        ["Dense-2 (Output)",      "Linear(32 → 1)",                                      "Final scalar prediction"],
    ],
    col_widths=[1.8, 3.3, 2.2],
)

heading2("5.4 Why GELU Activation?")
body("GELU (Gaussian Error Linear Unit) is used instead of ReLU in the Transformer feed-forward layers. "
     "GELU(x) = x · Φ(x) where Φ is the cumulative Gaussian distribution. It is smoother than ReLU "
     "(no hard zero cutoff) and has become standard in Transformer architectures (BERT, GPT).")

heading2("5.5 Learnable Positional Encoding")
body("Standard Transformers use fixed sinusoidal positional encodings. This implementation uses a "
     "learnable parameter tensor of shape (1, 100, 64) that is added to the projected input. "
     "This lets the model learn optimal positional representations specific to the 24-step solar "
     "cycle pattern, rather than using a generic fixed encoding.")

heading2("5.6 Training Configuration (PyTorch Lightning)")
add_table(
    ["Setting", "Value"],
    [
        ["Framework",              "PyTorch Lightning (L.LightningModule)"],
        ["Loss function",          "MSE (nn.functional.mse_loss)"],
        ["Optimizer",              "Adam, learning_rate=0.001"],
        ["Batch size",             "64"],
        ["Max epochs",             "100"],
        ["Validation split",       "10% of training data (last 10%, fixed seed)"],
        ["Early stopping",         "EarlyStopping(monitor=val_loss, patience=10)"],
        ["LR scheduler",           "ReduceLROnPlateau(factor=0.5, patience=5)"],
        ["Checkpointing",          "ModelCheckpoint: best val_loss saved as tft_model_best.ckpt"],
        ["Input normalization",    "scaler_X (StandardScaler on 14 features, fit on train)"],
        ["Output scaling",         "scaler_y for inverse_transform"],
        ["Saved as",               "tft_model.pt (state_dict) + tft_model_best.ckpt (full checkpoint)"],
    ],
    col_widths=[2.2, 4.1],
)

heading2("5.7 Performance & Observations")
performance_box("0.9959", "0.9472", "63.69", "28.58")
doc.add_paragraph("")
body("TFT shows the largest train/test gap of all models (0.9959 → 0.9472). This suggests overfitting "
     "despite dropout and early stopping. Possible reasons:")
bullet("The solar dataset is highly structured and periodic — tree models that explicitly encode "
       "lag features are better matched to this problem than attention mechanisms.")
bullet("The TFT architecture here is a simplified version without the full variable selection networks "
       "and interpretable gating of the original TFT paper (Lim et al., 2020).")
bullet("Despite lower standalone accuracy, the TFT contributes to ensemble robustness on days "
       "with unusual weather patterns where its attention-based global context helps.")

# ═══════════════════════════════════════════════════════════════════════════════
# MODEL 6 — TCN
# ═══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
heading1("Model 6: TCN (Temporal Convolutional Network)")

heading2("6.1 What Is a TCN?")
body("A Temporal Convolutional Network (TCN) applies dilated causal 1D convolutions to time series. "
     "It was proposed as an alternative to RNNs, offering several advantages:")
bullet("Parallelism: Unlike LSTMs, all timesteps are processed simultaneously.")
bullet("Stable gradients: No vanishing/exploding gradient problem (backprop through convolution is stable).")
bullet("Flexible receptive field: Dilated convolutions expand the effective receptive field exponentially "
       "without increasing parameter count.")

heading2("6.2 Causal Convolutions — No Future Leakage")
body("A standard Conv1D uses symmetric padding, meaning output at time t can see future values. "
     "Causal Conv1D pads only on the left (past side):")
centered_formula("padding  =  (kernel_size - 1) × dilation")
body("This padding is added to the left of the sequence, and the same amount is trimmed from the right "
     "of the output. This ensures the prediction at time t depends only on inputs at t, t-1, t-2, …")

heading2("6.3 Dilated Convolutions — Exponential Receptive Field")
body("In block i, dilation rate d_i = 2^i. With kernel_size=3:")
add_table(
    ["Block", "Dilation", "Effective Reception (timesteps covered)"],
    [
        ["Block 0", "1  (2^0)", "3 timesteps"],
        ["Block 1", "2  (2^1)", "7 timesteps"],
        ["Block 2", "4  (2^2)", "15 timesteps"],
        ["Block 3", "8  (2^3)", "29 timesteps  (covers full 24-step window + overlap)"],
    ],
    col_widths=[1.2, 1.5, 4.0],
)
doc.add_paragraph("")
body("With 4 blocks, the TCN can theoretically see the entire 24-step input window at block 3, "
     "without needing very deep networks or many parameters.")

heading2("6.4 Residual Blocks")
body("Each TemporalBlock has the structure:")
code_block("Input → CausalConv1d → BN → ReLU → Dropout(0.2)")
code_block("     → CausalConv1d → BN → ReLU → Dropout(0.2)")
code_block("     + Residual (1×1 Conv if channels differ) → ReLU → Output")
body("The residual (skip) connection adds the block input to its output. This:")
bullet("Enables very deep networks without degradation (like ResNets).")
bullet("Provides a gradient highway during backpropagation.")
bullet("Allows early layers to remain useful even as later layers specialize.")

heading2("6.5 Full Architecture")
add_table(
    ["Component", "Description", "Output Shape"],
    [
        ["Input (transposed)",  "x.transpose(1,2) — reorder to (batch, features, seq_len)", "(batch, 14, 24)"],
        ["Block 0",             "TemporalBlock(14→64, kernel=3, dilation=1, dropout=0.2)",  "(batch, 64, 24)"],
        ["Block 1",             "TemporalBlock(64→64, kernel=3, dilation=2, dropout=0.2)",  "(batch, 64, 24)"],
        ["Block 2",             "TemporalBlock(64→64, kernel=3, dilation=4, dropout=0.2)",  "(batch, 64, 24)"],
        ["Block 3",             "TemporalBlock(64→64, kernel=3, dilation=8, dropout=0.2)",  "(batch, 64, 24)"],
        ["Last timestep",       "out[:, :, -1] — final timestep only",                     "(batch, 64)"],
        ["LayerNorm",           "LayerNorm(64)",                                             "(batch, 64)"],
        ["Dense-1",             "Linear(64→32) + GELU + Dropout(0.2)",                     "(batch, 32)"],
        ["Output",              "Linear(32→1)",                                              "(batch, 1)"],
    ],
    col_widths=[1.6, 3.6, 1.6],
)

heading2("6.6 Training Configuration (PyTorch Lightning)")
body("Identical to TFT: PyTorch Lightning, Adam(lr=0.001), MSE loss, batch=64, max_epochs=100, "
     "EarlyStopping(patience=10), ModelCheckpoint. "
     "Saved as: tcn_model.pt + tcn_model_best.ckpt.")

heading2("6.7 TCN vs. CNN-LSTM — Key Differences")
add_table(
    ["Aspect", "CNN-LSTM (Keras)", "TCN (PyTorch)"],
    [
        ["Causality",         "Not strictly causal (LSTM handles ordering)", "Strictly causal (no future leakage at conv level)"],
        ["Dilation",          "None (kernel=3, no dilation)",                "Exponential: 1, 2, 4, 8"],
        ["Recurrence",        "Yes (BiLSTM layers)",                         "No recurrence — fully convolutional"],
        ["Receptive field",   "Limited by LSTM hidden size",                 "Explicitly controlled by dilation"],
        ["Parallelism",       "LSTM limits parallelism",                     "Fully parallel (all timesteps at once)"],
        ["Test R²",           "0.9896",                                      "0.9461"],
        ["Test RMSE",         "28.22 W/m²",                                  "64.38 W/m²"],
    ],
    col_widths=[1.8, 2.4, 2.6],
)

heading2("6.8 Performance")
performance_box("0.9936", "0.9461", "64.38", "28.83")
doc.add_paragraph("")
body("TCN and TFT show similar overfitting patterns, achieving near-identical test scores "
     "(0.9461 vs 0.9472). Both deep learning architectures from the tft/ package appear to need "
     "more training data or stronger regularization to match the tree models on this dataset.")

# ═══════════════════════════════════════════════════════════════════════════════
# COMPARATIVE SUMMARY
# ═══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
heading1("7. Comparative Summary & Model Selection Rationale")

heading2("7.1 Full Performance Comparison")
add_table(
    ["Model", "Framework", "Train R²", "Test R²", "Test RMSE", "Test MAE", "Train/Test Gap"],
    [
        ["XGBoost",      "xgboost",          "0.9981", "0.9933", "22.62", "9.69",  "0.0048 (minimal)"],
        ["Random Forest","scikit-learn",      "0.9963", "0.9933", "22.77", "9.88",  "0.0030 (minimal)"],
        ["CNN-LSTM",     "TensorFlow/Keras",  "0.9891", "0.9896", "28.22", "17.01", "−0.0005 (none)"],
        ["LSTM",         "TensorFlow/Keras",  "0.9876", "0.9873", "31.30", "20.19", "0.0003 (none)"],
        ["TFT",          "PyTorch Lightning", "0.9959", "0.9472", "63.69", "28.58", "0.0487 (moderate)"],
        ["TCN",          "PyTorch Lightning", "0.9936", "0.9461", "64.38", "28.83", "0.0475 (moderate)"],
    ],
    col_widths=[1.4, 1.5, 0.9, 0.9, 1.0, 0.9, 1.7],
)
doc.add_paragraph("")

heading2("7.2 Why Each Model Still Belongs in the Ensemble")
add_table(
    ["Model", "Unique Strength in Ensemble"],
    [
        ["XGBoost",       "Best raw accuracy; autoregressive lag features give explicit time memory; contributes most to ensemble (30% weight)"],
        ["Random Forest", "Bootstrapped diversity reduces variance; robust to individual outlier hours; stable backup to XGBoost"],
        ["CNN-LSTM",       "Captures smooth intra-day ramp shapes (dawn/dusk) that tree models may quantize into hard steps"],
        ["LSTM",           "Multi-day seasonal trend learning; beneficial on long-horizon winter days with unusual radiation profiles"],
        ["TFT",            "Global attention: can detect cross-hour correlations (e.g., morning fog predicting reduced afternoon radiation)"],
        ["TCN",            "Fully parallel; causal guarantee; receptive field explicitly tuned to 24-hour window via exponential dilation"],
    ],
    col_widths=[1.4, 5.0],
)
doc.add_paragraph("")

heading2("7.3 Common Training Infrastructure")
bullet("Framework: XGBoost/RF via scikit-learn API; LSTM/CNN-LSTM via TensorFlow/Keras; TFT/TCN via PyTorch Lightning.")
bullet("All models use RANDOM_STATE=42 for reproducibility.")
bullet("All scalers (StandardScaler) are fit exclusively on train data to prevent data leakage.")
bullet("All models evaluated on the same chronological 20% test set (2023–2025).")
bullet("Metrics: R², MAE, RMSE computed on original W/m² scale (inverse-transformed where needed).")

# ═══════════════════════════════════════════════════════════════════════════════
# SAVE
# ═══════════════════════════════════════════════════════════════════════════════
out_path = r"d:\Research Paper\Solar FYP (Faisal)\Project 1\Solar_FYP_Model_Explanations.docx"
doc.save(out_path)
print(f"Saved: {out_path}")
