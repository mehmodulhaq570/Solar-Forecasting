"""Generate Word report: All Model Explanations"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
section = doc.sections[0]
section.top_margin    = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin   = Inches(1.25)
section.right_margin  = Inches(1.25)

# ── Helpers ───────────────────────────────────────────────────────────────────
def heading1(text):
    p = doc.add_heading(text, level=1)
    p.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    return p

def heading2(text):
    p = doc.add_heading(text, level=2)
    p.runs[0].font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
    return p

def body(text):
    p = doc.add_paragraph(text)
    p.runs[0].font.size = Pt(11)
    return p

def bullet(text):
    p = doc.add_paragraph(text, style="List Bullet")
    for run in p.runs:
        run.font.size = Pt(11)
    return p

def add_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header row
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        run = hdr[i].paragraphs[0].runs[0]
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        # Blue fill
        tc = hdr[i]._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), '1F497D')
        tcPr.append(shd)
    # Data rows
    for ri, row_data in enumerate(rows):
        cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row_data):
            cells[ci].text = str(val)
            cells[ci].paragraphs[0].runs[0].font.size = Pt(10)
            if ri % 2 == 1:
                tc = cells[ci]._tc
                tcPr = tc.get_or_add_tcPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:color'), 'auto')
                shd.set(qn('w:fill'), 'DCE6F1')
                tcPr.append(shd)
    if col_widths:
        for ri, row in enumerate(table.rows):
            for ci, cell in enumerate(row.cells):
                cell.width = Inches(col_widths[ci])
    return table

# ═══════════════════════════════════════════════════════════════════════════════
# TITLE
# ═══════════════════════════════════════════════════════════════════════════════
title = doc.add_heading("Solar Radiation Forecasting System", 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

sub = doc.add_paragraph("Technical Report: Dataset, Preprocessing, Feature Engineering & Ensemble")
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.runs[0].bold = True
sub.runs[0].font.size = Pt(12)
sub.runs[0].font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

doc.add_paragraph("")

# ═══════════════════════════════════════════════════════════════════════════════
# 1. DATASET & PREPROCESSING
# ═══════════════════════════════════════════════════════════════════════════════
heading1("1. Dataset & Preprocessing")

heading2("1.1 Data Source")
body("The dataset is sourced from the NASA POWER (Prediction of Worldwide Energy Resources) API — "
     "hourly meteorological and solar radiation data for Lahore, Pakistan (31.56°N, 74.35°E).")
body("File: lahore_hourly_filled.csv")
bullet("Date range: January 2018 – October 2025")
bullet("Temporal resolution: 1 hour")
bullet("Approximate size: ~68,000 hourly records")

heading2("1.2 Raw Columns")
add_table(
    ["Column", "Description", "Unit"],
    [
        ["SolarRadiation",    "Target variable — Global Horizontal Irradiance (GHI)",  "W/m²"],
        ["ClearSkyRadiation", "Theoretical maximum radiation under clear sky",           "W/m²"],
        ["DirectRadiation",   "Direct Normal Irradiance (DNI)",                          "W/m²"],
        ["DiffuseRadiation",  "Diffuse Horizontal Irradiance (DHI)",                     "W/m²"],
        ["SolarZenith",       "Sun angle from vertical (≥90° = night)",                  "degrees"],
        ["Temperature",       "Air temperature at 2 m",                                  "°C"],
        ["HumiditySpecific",  "Specific humidity",                                       "g/kg"],
        ["HumidityRelative",  "Relative humidity",                                       "%"],
        ["Pressure",          "Surface pressure",                                        "kPa"],
        ["WindSpeed",         "Wind speed at 10 m",                                     "m/s"],
        ["WindDirection",     "Wind direction at 10 m",                                  "degrees"],
    ],
    col_widths=[1.6, 3.2, 1.0],
)
doc.add_paragraph("")

heading2("1.3 Preprocessing Steps")
body("The following steps are applied in training.py before any model training:")
bullet("Parse datetime index and sort chronologically.")
bullet("df.apply(pd.to_numeric, errors='coerce') — coerce all columns to numeric, handling any string artifacts from the CSV.")
bullet("Drop rows where SolarRadiation is NaN.")
bullet("Night hours are intentionally kept (not filtered) — models learn that SolarZenith ≥ 90° implies zero radiation.")
body("A derived feature is then computed:")
bullet("ClearnessIndex = SolarRadiation / ClearSkyRadiation, clipped to [0, 1].")
bullet("Zero-safe: when ClearSkyRadiation = 0, ClearnessIndex = 0 (avoids division by zero).")
bullet("Interpretation: values near 1 = clear sky; values near 0 = heavy cloud cover.")

# ═══════════════════════════════════════════════════════════════════════════════
# 2. TRAIN/TEST SPLIT & LEAKAGE PREVENTION
# ═══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
heading1("2. Train/Test Split & Leakage Prevention")

heading2("2.1 Chronological 80/20 Split")
body("A strict chronological split is used — no shuffling at any stage:")
bullet("Training set:  ~2018 → mid-2023  (80% of records)")
bullet("Test set:      ~mid-2023 → Oct 2025  (20% of records)")
body("Code: split_idx = int(len(df_day) * 0.8)")

heading2("2.2 Leakage Prevention Measures")
add_table(
    ["Measure", "How it is enforced"],
    [
        ["Scaler fitting",       "All StandardScalers are fit exclusively on training data, then transform() is applied to test data."],
        ["Lag feature safety",   "Lag features use .shift(lag) — only past values are used; no future values enter the feature window."],
        ["Sequence assignment",  "Sequences spanning the train/test boundary are assigned to the set that contains their target timestamp, not their window start."],
        ["Validation split",     "LSTM/CNN-LSTM use validation_split=0.1 on the training set — Keras takes the last 10% of training chronologically (no shuffle)."],
        ["Early stopping",       "EarlyStopping(patience=15, restore_best_weights=True) — stops on validation loss without touching test data."],
        ["Learning rate decay",  "ReduceLROnPlateau(patience=5) — adapts on validation, not test performance."],
    ],
    col_widths=[1.8, 4.5],
)

# ═══════════════════════════════════════════════════════════════════════════════
# 3. FEATURE ENGINEERING
# ═══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
heading1("3. Feature Engineering")

heading2("3.1 Tree Models (XGBoost, Random Forest) — 41 Features")
body("17 base features plus 24 autoregressive lag features:")
add_table(
    ["Category", "Features", "Count"],
    [
        ["Raw time",        "hour, month, day_of_year",                                                                                    "3"],
        ["Cyclical time",   "hour_sin, hour_cos (period 24)\nmonth_sin, month_cos (period 12)\ndoy_sin, doy_cos (period 365)",             "6"],
        ["Solar geometry",  "SolarZenith, ClearSkyRadiation",                                                                              "2"],
        ["Weather",         "Temperature, HumiditySpecific, HumidityRelative, Pressure, WindSpeed, WindDirection",                         "6"],
        ["Autoregressive",  "lag_1, lag_2, …, lag_24  (past 24 hours of SolarRadiation)",                                                 "24"],
    ],
    col_widths=[1.6, 4.0, 0.7],
)
doc.add_paragraph("")
body("Rationale for cyclical encoding: Raw hour=23 and hour=0 are numerically far apart but temporally "
     "adjacent. Sine/cosine encoding wraps this into a continuous circle, preventing the model from "
     "treating midnight as discontinuous.")
body("Rationale for SolarZenith: A physical constraint — when Zenith ≥ 90° the sun is below the horizon "
     "and radiation must be zero. This acts as the single most discriminative binary feature for night prediction.")

heading2("3.2 Sequence Models (LSTM, CNN-LSTM, TFT, TCN) — 15 Features × 24 Timesteps")
body("14 features (no raw time integers, no explicit lag columns) fed as a 24-step sliding window. "
     "The sequence window itself carries the temporal history equivalent to lag features.")
add_table(
    ["Category", "Features"],
    [
        ["Cyclical time",   "hour_sin, hour_cos, month_sin, month_cos, doy_sin, doy_cos"],
        ["Solar geometry",  "SolarZenith, ClearSkyRadiation"],
        ["Weather",         "Temperature, HumiditySpecific, HumidityRelative, Pressure, WindSpeed, WindDirection"],
        ["Target (in seq)", "SolarRadiation (past values included in input window)"],
    ],
    col_widths=[1.8, 4.5],
)
doc.add_paragraph("")
body("Input tensor shape: (N, 24, 15) — N samples × 24 timesteps × 15 features.")
body("All features are scaled with a StandardScaler fit on training data. A separate scaler_y is used "
     "for the target column to allow inverse_transform of predictions back to W/m².")

# ═══════════════════════════════════════════════════════════════════════════════
# 4. ENSEMBLE WEIGHT FORMATION
# ═══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
heading1("4. Ensemble Weight Formation")

heading2("4.1 Training-Derived Weights (training.py)")
body("After all models are trained, weights are assigned proportionally to their Test R² scores:")
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("w_i  =  max(0, R²_i)  /  Σ_j  max(0, R²_j)")
run.bold = True
run.font.size = Pt(12)
body("This ensures better-performing models receive higher influence, and any model with negative R² "
     "is excluded (weight floored to 0). The weights are saved to saved_models/ensemble_weights.pkl.")

heading2("4.2 Production Weights (frontend/config.py)")
body("Two manually tuned weight sets exist at runtime, depending on which external forecast API is used:")
add_table(
    ["Model", "NASA POWER Weight", "Open-Meteo Weight", "Rationale"],
    [
        ["XGBoost",      "0.30", "0.20", "NASA data format closely matches XGBoost training data; down-weighted when using real-time API"],
        ["Random Forest","0.25", "0.18", "Same reasoning as XGBoost"],
        ["LSTM",         "0.08", "0.12", "Sequence models valued more when real-time weather context is available"],
        ["CNN-LSTM",     "0.07", "0.10", "Same as LSTM"],
        ["TFT",          "0.18", "0.25", "Transformer attention handles uncertainty better with real-time data"],
        ["TCN",          "0.12", "0.15", "Causal convolutions adapt well to recent signal patterns"],
        ["Total",        "1.00", "1.00", "—"],
    ],
    col_widths=[1.4, 1.5, 1.5, 2.4],
)

heading2("4.3 Per-Hour Calibration Multipliers")
body("After the weighted blend, a per-hour scalar is applied to correct for systematic timing biases:")
bullet("NASA POWER calibration: Hours 0–6 and 17–23 → ×0.0 (hard zero for night). Hours 7–16 → ×1.0.")
bullet("Open-Meteo calibration: More granular ramp curve to counteract underestimation of the late-afternoon "
       "value by Open-Meteo. Hour 17 carries a ×10.0 multiplier, reflecting a known API bias at dusk.")
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Final Ensemble(h)  =  Cal(h)  ×  Σ_i  [ w_i  ×  ŷ_i(h) ]")
run.bold = True
run.font.size = Pt(12)

# ═══════════════════════════════════════════════════════════════════════════════
# 5. MODEL ACCURACY
# ═══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
heading1("5. Model Accuracy")

heading2("5.1 Individual Model Performance")
add_table(
    ["Model", "Train R²", "Test R²", "Test RMSE (W/m²)", "Test MAE (W/m²)"],
    [
        ["XGBoost",      "0.9981", "0.9933", "22.62", "9.69"],
        ["Random Forest","0.9963", "0.9933", "22.77", "9.88"],
        ["CNN-LSTM",     "0.9891", "0.9896", "28.22", "17.01"],
        ["LSTM",         "0.9876", "0.9873", "31.30", "20.19"],
        ["TFT",          "0.9959", "0.9472", "63.69", "28.58"],
        ["TCN",          "0.9936", "0.9461", "64.38", "28.83"],
    ],
    col_widths=[1.5, 1.2, 1.2, 1.7, 1.7],
)
doc.add_paragraph("")

heading2("5.2 Key Observations")
bullet("Tree models (XGBoost, Random Forest) dominate on test metrics with Test R² = 0.9933 and "
       "MAE below 10 W/m². This is why they receive the highest ensemble weights (0.25–0.30).")
bullet("CNN-LSTM outperforms plain LSTM: the Conv1D layers extract local temporal patterns before "
       "the BiLSTM layers process longer-range dependencies, yielding ~3 W/m² lower RMSE.")
bullet("TFT and TCN show overfitting: Train R² ~0.995 vs Test R² ~0.947. The gap indicates these "
       "models did not generalize as well on the 2023–2025 test window, possibly due to fewer training "
       "epochs or architectural under-regularization.")
bullet("Ensemble advantage: A single R² number cannot capture the robustness benefit. On extreme or "
       "out-of-distribution days (unusual cloud patterns, storm events), the ensemble hedges against "
       "catastrophic failure of any single model by blending diverse predictions.")

heading2("5.3 Why the Ensemble Is Still Preferred")
body("Even though XGBoost/RF alone achieve the highest individual Test R², the ensemble is used in "
     "production for several reasons:")
bullet("Variance reduction: Combining models with different architectures (tree, recurrent, convolutional, "
       "attention) reduces prediction variance on unseen dates outside the 2018–2025 training window.")
bullet("API calibration: The production ensemble integrates real-time API forecasts (NASA POWER / "
       "Open-Meteo) as a soft prior, correcting for weather anomalies that historical ML models "
       "cannot anticipate.")
bullet("Hour-level reliability: The calibration multipliers explicitly correct known systematic biases "
       "(night zeros, dusk underestimation) that individual models may still predict as small non-zero "
       "values due to noisy training data.")

# ═══════════════════════════════════════════════════════════════════════════════
# SAVE
# ═══════════════════════════════════════════════════════════════════════════════
out_path = r"d:\Research Paper\Solar FYP (Faisal)\Project 1\Solar_FYP_Technical_Report.docx"
doc.save(out_path)
print(f"Saved: {out_path}")
